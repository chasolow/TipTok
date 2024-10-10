import streamlit as st
import random
from datetime import datetime
from docx import Document

# Функция для заполнения шаблона
def fill_template(params):
    doc = Document("C:\\Users\\wanss\\OneDrive\\Рабочий стол\\TipTok\\Форма для запросов\\КЭЭ.docx")
    
    # Замена текста в абзацах
    for paragraph in doc.paragraphs:
        for key, value in params.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    # Замена текста в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in params.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    # Сохранение нового документа
    today = datetime.now()
    random_number = random.randint(10, 99)
    new_doc_name = f"KЭЭ_заполненный_KЭЭ-{today.strftime('%m-%d-%y')}-{random_number}.docx"
    new_doc_path = f"C:\\Users\\wanss\\OneDrive\\Рабочий стол\\TipTok\\Форма для запросов\\{new_doc_name}"
    doc.save(new_doc_path)
    return new_doc_path

# Основная функция
def main():
    st.title("Заполните форму для запроса")

    # Параметры формы
    params = {
        "01_Номер_письма": f"КЭЭ-{datetime.now().strftime('%m-%y')}/{random.randint(10, 99)}",
        "02_Дата_письма": datetime.now().strftime("%d.%m.%Y"),
        "03_Должность_адресата": st.selectbox("Должность адресата", ["Директор", "Главный инженер", "Заместитель генерального директора по транспорту электрической энергии"]),
        "04.1_Филиал_адресата": st.selectbox("Филиал адресата", ["Новосибирские городские электрические сети", "Приобские электрические сети", "Восточные электрические сети", "Черепановские электрические сети", "Западные электрические сети", "Чулымские электрические сети", "Карасукские электрические сети", "Татарские электрические сети"]),
        "04.2_Организация_адресата": st.selectbox("Организация адресата", ["АО \"РЭС\"", "ООО \"СЭТ54\"", "ООО \"ПСК\"", "Иная СО"]),
        "05_Фамилия_И.О._адресата": st.text_input("Фамилия И.О. адресата"),
        "06_email_адресата": st.text_input("Email адресата"),
        "07_Имя_Отчество_адресата": st.text_input("Имя Отчество адресата"),
        "08_Проектная_организация": st.text_input("Проектная организация"),
        "09_Объект": st.text_input("Объект"),
        "10_Техусловия": st.text_input("Технические условия"),
        "11_Шифр": st.text_input("Шифр"),
        "12_Должность_отправителя": st.text_input("Должность отправителя"),
        "13_И.О._Фамилия_отправителя": st.text_input("И.О. Фамилия отправителя"),
        "14_Фамилия_И.О._исполнителя": st.text_input("Фамилия И.О. исполнителя"),
        "15_email_исполнителя": st.text_input("Email исполнителя"),
        "16_телефон_исполнителя": st.text_input("Телефон исполнителя"),
    }

    if st.button("Сохранить документ"):
        output_path = fill_template(params)
        st.success(f"Документ сохранён: {output_path}")

if __name__ == "__main__":
    main()
